#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para generar un documento Word con una frase repetida mil veces
"""

from docx import Document
from docx.shared import Inches
import os

def generar_documento_con_frase(frase, num_repeticiones=1000, nombre_archivo="documento_frase.docx"):
    """
    Genera un documento Word con una frase repetida el número de veces especificado
    
    Args:
        frase (str): La frase a repetir
        num_repeticiones (int): Número de veces a repetir la frase (por defecto 1000)
        nombre_archivo (str): Nombre del archivo de salida
    """
    try:
        # Crear un nuevo documento
        doc = Document()
        
        # Agregar título
        titulo = doc.add_heading(f'Documento con frase repetida {num_repeticiones} veces', 0)
        
        # Agregar información del documento
        doc.add_paragraph(f'Frase: "{frase}"')
        doc.add_paragraph(f'Número de repeticiones: {num_repeticiones}')
        doc.add_paragraph('')  # Línea en blanco
        
        # Agregar la frase repetida
        for i in range(num_repeticiones):
            # Agregar número de línea para mejor organización
            parrafo = doc.add_paragraph(f'{i+1:04d}. {frase}')
            
            # Agregar un salto de línea cada 10 frases para mejor legibilidad
            if (i + 1) % 10 == 0:
                doc.add_paragraph('')  # Línea en blanco
        
        # Guardar el documento
        doc.save(nombre_archivo)
        
        print(f"✅ Documento generado exitosamente: {nombre_archivo}")
        print(f"📄 Ubicación: {os.path.abspath(nombre_archivo)}")
        print(f"📊 Total de frases escritas: {num_repeticiones}")
        
        return True
        
    except Exception as e:
        print(f"❌ Error al generar el documento: {str(e)}")
        return False

def main():
    """Función principal del script"""
    print("=" * 60)
    print("🖊️  GENERADOR DE DOCUMENTO WORD CON FRASE REPETIDA")
    print("=" * 60)
    
    # Solicitar datos al usuario
    print("\n📝 Ingresa la frase que deseas repetir:")
    frase = input("Frase: ").strip()
    
    if not frase:
        print("❌ Error: La frase no puede estar vacía")
        return
    
    print(f"\n🔢 ¿Cuántas veces quieres repetir la frase? (por defecto: 1000)")
    try:
        num_repeticiones = input("Número de repeticiones: ").strip()
        if num_repeticiones:
            num_repeticiones = int(num_repeticiones)
            if num_repeticiones <= 0:
                print("❌ Error: El número de repeticiones debe ser mayor a 0")
                return
        else:
            num_repeticiones = 1000
    except ValueError:
        print("❌ Error: Ingresa un número válido")
        return
    
    print(f"\n📁 ¿Qué nombre quieres darle al archivo? (por defecto: documento_frase.docx)")
    nombre_archivo = input("Nombre del archivo: ").strip()
    if not nombre_archivo:
        nombre_archivo = "documento_frase.docx"
    elif not nombre_archivo.endswith('.docx'):
        nombre_archivo += '.docx'
    
    # Confirmar antes de generar
    print(f"\n📋 Resumen:")
    print(f"   • Frase: '{frase}'")
    print(f"   • Repeticiones: {num_repeticiones}")
    print(f"   • Archivo: {nombre_archivo}")
    
    confirmacion = input("\n¿Proceder con la generación? (s/n): ").lower().strip()
    if confirmacion not in ['s', 'si', 'sí', 'y', 'yes']:
        print("❌ Operación cancelada")
        return
    
    # Generar el documento
    print(f"\n🔄 Generando documento...")
    if generar_documento_con_frase(frase, num_repeticiones, nombre_archivo):
        print(f"\n🎉 ¡Documento generado exitosamente!")
        print(f"📂 Puedes abrir el archivo: {nombre_archivo}")
    else:
        print(f"\n💥 Error al generar el documento")

if __name__ == "__main__":
    main() 